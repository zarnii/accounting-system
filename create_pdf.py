import docx
#from db import DataBase

def create_pdf(data):
	# create Document
	doc = docx.Document()


	table = doc.add_table(rows = 1, cols = 4, style="Table Grid")
	all_price = 0

	row = table.rows[0].cells
	row[0].text = 'id'
	row[1].text = 'название'
	row[2].text = 'количесво'
	row[3].text = 'цена'

	for i, name, count, price in data:
		# Adding a row and then adding data in it.
		row = table.add_row().cells
		# Converting id to string as table can only take string input
		row[0].text = str(i)
		row[1].text = str(name)
		row[2].text = str(count)
		row[3].text = str(price*count)
		all_price += price*count

	doc.add_paragraph(f'Всего: {all_price}')
	

	# store
	doc.save('pdf/output.docx')


if __name__ == "__main__":
	create_pdf()